import os, re
from flask import Flask, request, jsonify, render_template, Response
from flask_cors import CORS
import PyPDF2
from docx import Document
from weasyprint import HTML as PDFGen
from pdf_parser import ResumeParser, ResumeParser_OpenAI  # wherever you placed the class
import utils  # <-- import your util module
from docx import Document
from docx.shared import Pt
import io
from bs4 import BeautifulSoup
app = Flask(__name__)
CORS(app)


# In-memory job description
job_description_data = ""

KNOWN_HEADINGS = ['Technical Skills','Experience','Education','Projects','Publications']
# instantiate once
parser = ResumeParser(known_headings=KNOWN_HEADINGS)
parser2 = ResumeParser_OpenAI()
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# In‐memory resume state
sections_data = []
header_html = ''

# Preloaded suggestions JSON
suggestions_data = {
  "existing": [
    { "section": "Technical Skills", "subsection": "Languages & databases", "category": "Technical Skills",
      "keywords": ["python"],
      "suggestions": ["Consider highlighting Python as your primary programming language tailored for machine learning and NLP applications."] },
    { "section": "Technical Skills", "subsection": "Libraries", "category": "Technical Skills",
      "keywords": ["Tensorflow","Keras","PyTorch","Scikit-learn","NLP toolkit"],
      "suggestions": ["Emphasize experience with TensorFlow, Keras, and PyTorch by detailing specific projects or models you've worked on."] },
    { "section": "Technical Skills", "subsection": "Libraries", "category": "Technical Skills",
      "keywords": ["Natural Language Processing (NLP)"],
      "suggestions": ["Expand on specific NLP techniques you've implemented or studied, particularly in chatbot-related projects."] },
    { "section": "Experience", "subsection": "Medical Aid", "category": "Experience",
      "keywords": ["ChatBot","NLP models"],
      "suggestions": ["Highlight your contribution in developing the ChatBot with a focus on NLP applications to enhance usability."] },
    { "section": "Experience", "subsection": "Old Dominion University", "category": "Experience",
      "keywords": ["Machine Learning","Deep Learning"],
      "suggestions": ["Detail your research efforts specific to ML and DL models, especially in classification problems."] }
  ],
  "missing": [
    { "section": "Technical Skills", "subsection": "Technical Skills", "category": "Technical Skills",
      "keywords": ["Generative AI","APIs"],
      "suggestions": ["Include your experience with generative AI techniques and any APIs used for model deployment."] },
    { "section": "Technical Skills", "subsection": "Technical Skills", "category": "Technical Skills",
      "keywords": ["Front-end and Back-end integration"],
      "suggestions": ["Add experience related to front-end and back-end connectivity, emphasizing any integration projects you've worked on."] },
    { "section": "Experience", "subsection": "Experience", "category": "Experience",
      "keywords": ["Data sources","cross-functional collaboration"],
      "suggestions": ["Include details about working with disparate data sources and collaborating with cross-functional teams to align projects."] },
    { "section": "General Skills", "subsection": "General Skills", "category": "Soft Skills",
      "keywords": ["communication skills","problem-solving","analytical","organizational"],
      "suggestions": ["Highlight any relevant experiences that reflect strong communication and problem-solving skills, especially in technical contexts."] }
  ]
}



def split_by_main_headings(lines):
    blocks, curr, buf = [], None, []
    for raw in lines:
        t = raw.strip()
        if t.lower() in [h.lower() for h in KNOWN_HEADINGS]:
            if curr: blocks.append((curr, buf))
            curr = next(h for h in KNOWN_HEADINGS if h.lower()==t.lower())
            buf = []
        else:
            if curr: buf.append(raw)
    if curr: blocks.append((curr, buf))
    return blocks

def split_into_subsections(lines):
    subs, title, buf = [], None, []
    for raw in lines:
        if '|' in raw:
            if title: subs.append((title, buf))
            title, buf = raw.strip(), []
        else:
            buf.append(raw)
    if title: subs.append((title, buf))
    return subs or [(None, lines)]

def build_html_from_lines(lines):
    html_parts = []
    in_list = False
    for raw in lines:
        s = raw.strip()
        if s.startswith(('•','-','*','–')):
            if not in_list:
                html_parts.append('<ul>')
                in_list = True
            item = s.lstrip('•-*– ').strip()
            html_parts.append(f'<li>{item}</li>')
        else:
            if in_list:
                html_parts.append('</ul>')
                in_list = False
            if s:
                html_parts.append(f'<p>{s}</p>')
    if in_list:
        html_parts.append('</ul>')
    return ''.join(html_parts)

def parse_pdf_lines(path):
    reader = PyPDF2.PdfReader(path)
    out = []
    for page in reader.pages:
        txt = page.extract_text()
        if txt: out.extend(txt.splitlines())
    return out

def parse_docx_lines(path):
    doc = Document(path)
    return [p.text for p in doc.paragraphs if p.text.strip()]

@app.route('/')
def index():
    return render_template('index.html')

# @app.route('/upload', methods=['POST'])
# def upload():
#     global sections_data, header_html
#     f = request.files.get('file')
#     if not f: return jsonify({'error':'No file provided'}),400
#     _, ext = os.path.splitext(f.filename.lower())
#     if ext not in ('.pdf','.doc','.docx','.doc'):
#         return jsonify({'error':'Bad file type'}),400
#     fp = os.path.join(UPLOAD_FOLDER, f.filename)
#     f.save(fp)
#     raw = parse_pdf_lines(fp) if ext=='.pdf' else parse_docx_lines(fp)

#     # extract header
#     idx = next((i for i,line in enumerate(raw)
#                 if line.strip().lower() in [h.lower() for h in KNOWN_HEADINGS]), len(raw))
#     header_html = build_html_from_lines(raw[:idx])

#     # build sections
#     secs, cnt = [], 0
#     for heading, block in split_by_main_headings(raw):
#         for sub_title, body in split_into_subsections(block):
#             subtitle = sub_title or heading
#             section_html = build_html_from_lines(body)
#             secs.append({'id':cnt,'section':heading,'subtitle':subtitle,'html':section_html})
#             cnt += 1
#     sections_data = secs
#     import pdb; pdb.set_trace()
#     return jsonify({'sections': sections_data})



@app.route('/upload', methods=['POST'])
def upload():
    global sections_data, header_html, resume_text

    f = request.files.get('file')
    if not f:
        return jsonify({'error':'No file provided'}), 400

    name, ext = os.path.splitext(f.filename.lower())
    if ext not in ('.pdf','.doc','.docx','.doc'):
        return jsonify({'error':'Bad file type'}), 400

    save_path = os.path.join(UPLOAD_FOLDER, f.filename)
    f.save(save_path)

    # delegate parsing
    header_html, sections_data, sections_pydantic, resume_text = parser2.parse(save_path)

    return jsonify({'sections': sections_data})


@app.route('/sections', methods=['GET'])
def list_sections():
    return jsonify(sections_data)

@app.route('/sections', methods=['POST'])
def create_section():
    global sections_data
    d = request.get_json() or {}
    new_id = max((s['id'] for s in sections_data), default=-1) + 1
    sec = {'id':new_id, 'section':d.get('section'), 'subtitle':d.get('subtitle'), 'html':d.get('html','')}
    sections_data.append(sec)
    return jsonify(sec),201

@app.route('/sections/<int:sec_id>', methods=['PUT'])
def update_section(sec_id):
    d = request.get_json() or {}
    for s in sections_data:
        if s['id'] == sec_id:
            s['subtitle'] = d.get('subtitle', s['subtitle'])
            s['html']     = d.get('html', s['html'])
            return jsonify(s)
    return jsonify({'error':'Not found'}),404

@app.route('/sections/<int:sec_id>', methods=['DELETE'])
def delete_section(sec_id):
    global sections_data
    for i, s in enumerate(sections_data):
        if s['id'] == sec_id:
            sections_data.pop(i)
            return '',204
    return jsonify({'error':'Not found'}),404

@app.route('/suggestions', methods=['GET'])
def get_suggestions():
    return jsonify(suggestions_data)

@app.route('/suggestions', methods=['POST'])
def post_suggestions():
    global suggestions_data
    d = request.get_json() or {}
    suggestions_data['existing'] = d.get('existing', [])
    suggestions_data['missing']  = d.get('missing', [])
    return jsonify(suggestions_data)

@app.route('/jobdesc', methods=['GET'])
def get_jobdesc():
    return jsonify({'job_description': job_description_data})

@app.route('/jobdesc', methods=['POST'])
def post_jobdesc():
    global job_description_data
    d = request.get_json() or {}
    job_description_data = d.get('job_description', '')
    return jsonify({'job_description': job_description_data})

# --- NEW: Similarity endpoint ---
@app.route('/similarity', methods=['POST'])
def similarity():
    payload = request.get_json() or {}
    resume_text = payload.get('resume', '')
    job_desc    = payload.get('job_description', '')
    # import pdb; pdb.set_trace()
    # print('job_desc: ', 'None' if (job_desc is None) or (job_desc is '') else job_desc)
    
    hybrid, semantic, keyword = utils.hybrid_resume_job_match_score(resume_text, job_desc)
    return jsonify({
        'hybrid_score': hybrid,
        'semantic_score': semantic,
        'keyword_score': keyword
    })

@app.route('/export', methods=['POST'])
def export_pdf():
    data = request.get_json() or {}
    sent = data.get('sections', [])
    by_main = {}
    for s in sent:
        by_main.setdefault(s['section'], []).append(s)
    rendered = render_template('export.html', header_html=header_html, sections_by_main=by_main)
    pdf = PDFGen(string=rendered).write_pdf()
    return Response(pdf, mimetype='application/pdf',
                    headers={'Content-Disposition':'attachment; filename=resume.pdf'})


# @app.route('/export-docx', methods=['POST'])
# @app.route('/export', methods=['POST'])
# def export_docx():
#     data = request.get_json() or {}
#     sent = data.get('sections', [])

#     # Group by section name
#     by_main = {}
#     for s in sent:
#         by_main.setdefault(s['section'], []).append(s)

#     # Create DOCX
#     doc = Document()
#     doc.styles['Normal'].font.name = 'Arial'
#     doc.styles['Normal'].font.size = Pt(11)

#     for section, entries in by_main.items():
#         # Section title
#         doc.add_heading(section, level=1)

#         for entry in entries:
#             subtitle = entry.get('subtitle', '')
#             html_block = entry.get('html', '')

#             # Add subtitle
#             if subtitle:
#                 doc.add_paragraph(subtitle, style='Intense Quote')

#             # Naive HTML to text conversion (only <p> and <ul><li> handled)
#             # lines = html_block.replace('</p>', '\n').replace('<p>', '').replace('</ul>', '').replace('<ul>', '').split('<li>')
#             soup = BeautifulSoup(html_block, 'html.parser')
#             text_lines = [line.strip() for line in soup.stripped_strings if line.strip()]
#             for line in text_lines:
#                 doc.add_paragraph(line, style='Normal')
#             # for line in lines:
#             #     clean = line.replace('</li>', '').strip()
#             #     if clean:
#             #         doc.add_paragraph(clean, style='List Bullet' if '<li>' in html_block else 'Normal')

#     # Return as a response
#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)

#     return Response(buffer.read(), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
#                     headers={'Content-Disposition': 'attachment; filename=resume.docx'})


@app.route("/preview", methods=["POST"])
def preview_resume():
    data = request.get_json()
    raw_sections = data.get("sections", [])
    # import pdb; pdb.set_trace()
    # Convert flat format to structured format expected by the template
    by_section = {}
    for s in raw_sections:
        section = s['section']
        s['html'] = s['html'].replace('<li>','')
        s['html'] = s['html'].replace('</li>','')
        s['html'] = s['html'].replace('<p>','')
        s['html'] = s['html'].replace('</p>','')
        s['html'] = s['html'].replace('<ul>','')
        s['html'] = s['html'].replace('</ul>','')
        if section not in by_section:
            by_section[section] = []
        by_section[section].append({
            'title': '',  # or extract from HTML if you wish
            'subtitle': s.get('subtitle', ''),
            'start_date': '',
            'end_date': '',
            'details': [line.strip() for line in s.get('html', '')
                        .replace('</p>', '')
                        .split('<p>') if line.strip()]
        })

    structured_sections = [
        {'section_name': key, 'items': items}
        for key, items in by_section.items()
    ]

    return render_template(
        'resume_template.html',
        name='Mory Gharasuie',
        location='Norfolk, VA',
        email='mory@example.com',
        phone='123-456-7890',
        linkedin='linkedin.com/in/mory',
        github='github.com/mory',
        sections=structured_sections  # ✅ fix here
    )
    
if __name__=='__main__':
    app.run(debug=True)
