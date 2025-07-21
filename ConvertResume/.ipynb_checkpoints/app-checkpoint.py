# import os, re
# from flask import Flask, request, jsonify, render_template, Response
# from flask_cors import CORS
# import PyPDF2
# from docx import Document
# from weasyprint import HTML as PDFGen
# from pdf_parser import ResumeParser_OpenAI  # wherever you placed the class
# import utils  # <-- import your util module
# # Define section/item types if needed
# from werkzeug.utils import secure_filename
# from collections import namedtuple

# ResumeItem = namedtuple("ResumeItem", ["title", "subtitle", "start_date", "end_date", "details", "extra"])
# ResumeSection = namedtuple("ResumeSection", ["section_name", "items"])

# app = Flask(__name__)
# CORS(app)


# # In-memory job description
# job_description_data = ""

# # KNOWN_HEADINGS = ['Technical Skills','Experience','Education','Projects','Publications']
# # instantiate once
# parser = ResumeParser_OpenAI()#ResumeParser(known_headings=KNOWN_HEADINGS)
# UPLOAD_FOLDER = 'uploads'
# os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# # In‐memory resume state
# sections_data = []
# header_html = ''


# # Preloaded suggestions JSON
# suggestions_data = {
#   "existing": [
#     { "section": "Technical Skills", "subsection": "Languages & databases", "category": "Technical Skills",
#       "keywords": ["python"],
#       "suggestions": ["Consider highlighting Python as your primary programming language tailored for machine learning and NLP applications."] },
#     { "section": "Technical Skills", "subsection": "Libraries", "category": "Technical Skills",
#       "keywords": ["Tensorflow","Keras","PyTorch","Scikit-learn","NLP toolkit"],
#       "suggestions": ["Emphasize experience with TensorFlow, Keras, and PyTorch by detailing specific projects or models you've worked on."] },
#     { "section": "Technical Skills", "subsection": "Libraries", "category": "Technical Skills",
#       "keywords": ["Natural Language Processing (NLP)"],
#       "suggestions": ["Expand on specific NLP techniques you've implemented or studied, particularly in chatbot-related projects."] },
#     { "section": "Experience", "subsection": "Medical Aid", "category": "Experience",
#       "keywords": ["ChatBot","NLP models"],
#       "suggestions": ["Highlight your contribution in developing the ChatBot with a focus on NLP applications to enhance usability."] },
#     { "section": "Experience", "subsection": "Old Dominion University", "category": "Experience",
#       "keywords": ["Machine Learning","Deep Learning"],
#       "suggestions": ["Detail your research efforts specific to ML and DL models, especially in classification problems."] }
#   ],
#   "missing": [
#     { "section": "Technical Skills", "subsection": "Technical Skills", "category": "Technical Skills",
#       "keywords": ["Generative AI","APIs"],
#       "suggestions": ["Include your experience with generative AI techniques and any APIs used for model deployment."] },
#     { "section": "Technical Skills", "subsection": "Technical Skills", "category": "Technical Skills",
#       "keywords": ["Front-end and Back-end integration"],
#       "suggestions": ["Add experience related to front-end and back-end connectivity, emphasizing any integration projects you've worked on."] },
#     { "section": "Experience", "subsection": "Experience", "category": "Experience",
#       "keywords": ["Data sources","cross-functional collaboration"],
#       "suggestions": ["Include details about working with disparate data sources and collaborating with cross-functional teams to align projects."] },
#     { "section": "General Skills", "subsection": "General Skills", "category": "Soft Skills",
#       "keywords": ["communication skills","problem-solving","analytical","organizational"],
#       "suggestions": ["Highlight any relevant experiences that reflect strong communication and problem-solving skills, especially in technical contexts."] }
#   ]
# }

# # -----------------------------------------------------------
# def generate_html_from_item(item: ResumeItem) -> str:
#     html = ""
#     if item.subtitle:
#         html += f"<p><em>{item.subtitle}</em></p>"
#     if item.start_date or item.end_date:
#         html += f"<p>{item.start_date or ''} – {item.end_date or ''}</p>"
#     if item.details:
#         html += "<ul>"
#         for detail in item.details:
#             if isinstance(detail, str) and detail.startswith("http"):
#                 html += f'<li><a href="{detail}" target="_blank">{detail}</a></li>'
#             else:
#                 html += f"<li>{detail}</li>"
#         html += "</ul>"
#     return html

# def serialize_sections(sections):
#     data = []
#     id_counter = 1
#     for section in sections:
#         for item in section.items:
#             data.append({
#                 "id": id_counter,
#                 "section": section.section_name,
#                 "subtitle": item.title,
#                 "html": generate_html_from_item(item)
#             })
#             id_counter += 1
#     return data

# @app.route("/sections")
# def serve_sections():
#     global sections_data
#     print("in serve_sections- sections_data: ", sections_data)
#     return jsonify(serialize_sections(sections_data))
# #-----------------------------------------------------------

# def split_by_main_headings(lines):
#     blocks, curr, buf = [], None, []
#     for raw in lines:
#         t = raw.strip()
#         if t.lower() in [h.lower() for h in KNOWN_HEADINGS]:
#             if curr: blocks.append((curr, buf))
#             curr = next(h for h in KNOWN_HEADINGS if h.lower()==t.lower())
#             buf = []
#         else:
#             if curr: buf.append(raw)
#     if curr: blocks.append((curr, buf))
#     return blocks

# def split_into_subsections(lines):
#     subs, title, buf = [], None, []
#     for raw in lines:
#         if '|' in raw:
#             if title: subs.append((title, buf))
#             title, buf = raw.strip(), []
#         else:
#             buf.append(raw)
#     if title: subs.append((title, buf))
#     return subs or [(None, lines)]

# def build_html_from_lines(lines):
#     html_parts = []
#     in_list = False
#     for raw in lines:
#         s = raw.strip()
#         if s.startswith(('•','-','*','–')):
#             if not in_list:
#                 html_parts.append('<ul>')
#                 in_list = True
#             item = s.lstrip('•-*– ').strip()
#             html_parts.append(f'<li>{item}</li>')
#         else:
#             if in_list:
#                 html_parts.append('</ul>')
#                 in_list = False
#             if s:
#                 html_parts.append(f'<p>{s}</p>')
#     if in_list:
#         html_parts.append('</ul>')
#     return ''.join(html_parts)

# def parse_pdf_lines(path):
#     reader = PyPDF2.PdfReader(path)
#     out = []
#     for page in reader.pages:
#         txt = page.extract_text()
#         if txt: out.extend(txt.splitlines())
#     return out

# def parse_docx_lines(path):
#     doc = Document(path)
#     return [p.text for p in doc.paragraphs if p.text.strip()]

# @app.route('/')
# def index():
#     return render_template('index.html')

# # @app.route('/upload', methods=['POST'])
# # def upload():
# #     global sections_data, header_html
# #     f = request.files.get('file')
# #     if not f: return jsonify({'error':'No file provided'}),400
# #     _, ext = os.path.splitext(f.filename.lower())
# #     if ext not in ('.pdf','.doc','.docx','.doc'):
# #         return jsonify({'error':'Bad file type'}),400
# #     fp = os.path.join(UPLOAD_FOLDER, f.filename)
# #     f.save(fp)
# #     raw = parse_pdf_lines(fp) if ext=='.pdf' else parse_docx_lines(fp)

# #     # extract header
# #     idx = next((i for i,line in enumerate(raw)
# #                 if line.strip().lower() in [h.lower() for h in KNOWN_HEADINGS]), len(raw))
# #     header_html = build_html_from_lines(raw[:idx])

# #     # build sections
# #     secs, cnt = [], 0
# #     for heading, block in split_by_main_headings(raw):
# #         for sub_title, body in split_into_subsections(block):
# #             subtitle = sub_title or heading
# #             section_html = build_html_from_lines(body)
# #             secs.append({'id':cnt,'section':heading,'subtitle':subtitle,'html':section_html})
# #             cnt += 1
# #     sections_data = secs
# #     return jsonify({'sections': sections_data})


    

# @app.route('/upload', methods=['POST'])
# def upload():
#     global sections_data, header_html, resume_text

#     # f = request.files.get('file')
#     # if not f:
#     #     return jsonify({'error':'No file provided'}), 400

#     # name, ext = os.path.splitext(f.filename.lower())
#     # if ext not in ('.pdf','.doc','.docx','.doc'):
#     #     return jsonify({'error':'Bad file type'}), 400

#     # save_path = os.path.join(UPLOAD_FOLDER, f.filename)
#     # f.save(save_path)

#     # # delegate parsing
#     # header_html, sections_data, resume_text = parser.parse(save_path)

#     # return jsonify({'sections': sections_data})
#     global sections
#     file = request.files["file"]
#     if not file:
#         return jsonify({"error": "No file uploaded"}), 400
#     filename = secure_filename(file.filename)
#     filepath = os.path.join(UPLOAD_FOLDER, filename)
#     file.save(filepath)

#     # Parse file to sections
#     try:
#         header_html, sections_data, resume_text = parser.parse(filepath)  # returns List[ResumeSection]
#         print("in upload method inside app.py")
#         return jsonify({"sections": serialize_sections(sections_data)})
#     except Exception as e:
#         print(e)
#         return jsonify({"error": str(e)}), 500


# @app.route('/sections', methods=['GET'])
# def list_sections():
#     # return jsonify(sections_data)
#     return serve_sections(section_data)
    

# @app.route('/sections', methods=['POST'])
# def create_section():
#     global sections_data
#     d = request.get_json() or {}
#     new_id = max((s['id'] for s in sections_data), default=-1) + 1
#     sec = {'id':new_id, 'section':d.get('section'), 'subtitle':d.get('subtitle'), 'html':d.get('html','')}
#     sections_data.append(sec)
#     return jsonify(sec),201

# @app.route('/sections/<int:sec_id>', methods=['PUT'])
# def update_section(sec_id):
#     d = request.get_json() or {}
#     for s in sections_data:
#         if s['id'] == sec_id:
#             s['subtitle'] = d.get('subtitle', s['subtitle'])
#             s['html']     = d.get('html', s['html'])
#             return jsonify(s)
#     return jsonify({'error':'Not found'}),404

# @app.route('/sections/<int:sec_id>', methods=['DELETE'])
# def delete_section(sec_id):
#     global sections_data
#     for i, s in enumerate(sections_data):
#         if s['id'] == sec_id:
#             sections_data.pop(i)
#             return '',204
#     return jsonify({'error':'Not found'}),404

# @app.route('/suggestions', methods=['GET'])
# def get_suggestions():
#     return jsonify(suggestions_data)

# @app.route('/suggestions', methods=['POST'])
# def post_suggestions():
#     global suggestions_data
#     d = request.get_json() or {}
#     suggestions_data['existing'] = d.get('existing', [])
#     suggestions_data['missing']  = d.get('missing', [])
#     return jsonify(suggestions_data)

# @app.route('/jobdesc', methods=['GET'])
# def get_jobdesc():
#     return jsonify({'job_description': job_description_data})

# @app.route('/jobdesc', methods=['POST'])
# def post_jobdesc():
#     global job_description_data
#     d = request.get_json() or {}
#     job_description_data = d.get('job_description', '')
#     return jsonify({'job_description': job_description_data})

# # --- NEW: Similarity endpoint ---
# @app.route('/similarity', methods=['POST'])
# def similarity():
#     payload = request.get_json() or {}
#     resume_text = payload.get('resume', '')
#     job_desc    = payload.get('job_description', '')
#     hybrid, semantic, keyword = utils.hybrid_resume_job_match_score(resume_text, job_desc)
#     return jsonify({
#         'hybrid_score': hybrid,
#         'semantic_score': semantic,
#         'keyword_score': keyword
#     })

# @app.route('/export', methods=['POST'])
# def export_pdf():
#     data = request.get_json() or {}
#     sent = data.get('sections', [])
#     by_main = {}
#     for s in sent:
#         by_main.setdefault(s['section'], []).append(s)
#     rendered = render_template('export.html', header_html=header_html, sections_by_main=by_main)
#     pdf = PDFGen(string=rendered).write_pdf()
#     return Response(pdf, mimetype='application/pdf',
#                     headers={'Content-Disposition':'attachment; filename=resume.pdf'})

# if __name__=='__main__':
#     app.run(debug=True)



from flask import Flask, render_template, request
from werkzeug.utils import secure_filename
# from werkzeug.serving import WSGIRequestHandler
# from werkzeug.serving import WSGIRequestHandler
import os
from celery import Celery
# from parser import parser  # This should return parsed Pydantic sections
import PyPDF2
from docx import Document
from flask_socketio import SocketIO, emit
from pdf_parser import ResumeParser_OpenAI  # wherever you placed the class
parser = ResumeParser_OpenAI()#ResumeParser(known_headings=KNOWN_HEADINGS)
app = Flask(__name__)
socketio = SocketIO(app, cors_allowed_origins="*")
# app.config['CELERY_BROKER_URL'] = 'redis://localhost:5000/0'
# celery = Celery(app.name, broker=app.config['CELERY_BROKER_URL'])
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
def extract_raw_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext == ".docx":
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        return "Unsupported file type."

@app.route("/", methods=["GET", "POST"])
def index():
    sections = None
    raw_text = ""
    filename = ""

    if request.method == "POST":
        file = request.files["resume"]
        if file and file.filename:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            raw_text = extract_raw_text(filepath)
            # parser_output = executor.submit(parser.parse(filepath))
            _, sections, _ = parser.parse(filepath)
            # _, sections, _ = parser_output.result()

    return render_template("index.html", sections=sections, raw_text=raw_text, filename=filename)

if __name__ == "__main__":
    # WSGIRequestHandler.timeout=120000
    app.run(debug=True)
